"""
Medical documents management endpoints for uploading, processing, and searching medical literature.
"""

from typing import List, Optional, Dict, Any, Union
from uuid import UUID
from fastapi import APIRouter, HTTPException, Depends, status, UploadFile, File, BackgroundTasks, Form, Query, Body
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from sqlalchemy import func, and_, or_, desc, asc
from loguru import logger

from app.core.database import get_db
from app.models.schemas import (
    DocumentUpload, DocumentResponse, DocumentSearch,
    PaginatedResponse, PaginationParams
)
from app.models.models import MedicalDocument, DocumentChunk, DocumentType, DocumentStatus, User
from app.schemas.document import DocumentStatistics
from app.core.config import settings
from app.services.document_service import document_service
import os
from datetime import datetime, timedelta, timezone
import json
from app.services.vector_service import vector_service
from app.services.ai_service import ai_service
from app.dependencies.auth import get_current_user, get_admin_user


router = APIRouter()


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    document_data: DocumentUpload = Depends(),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> DocumentResponse:
    """
    Upload and process a medical document.
    
    Args:
        background_tasks: Background tasks manager
        file: Uploaded document file
        document_data: Document metadata
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        Accepted document information
    """
    try:
        # Check user's document quota
        from app.core.config import settings
        
        # Admin users have no quota
        if not current_user.is_admin:
            doc_count = db.query(MedicalDocument).filter(
                MedicalDocument.uploaded_by == current_user.id
            ).count()
            
            if doc_count >= settings.max_documents_per_user:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail=f"Maximum document limit of {settings.max_documents_per_user} reached"
                )
        
        # Validate file type
        if not any(file.filename.lower().endswith(ext) for ext in settings.allowed_file_types):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not allowed. Allowed types: {settings.allowed_file_types}"
            )
        
        # Read file content
        content = await file.read()
        
        # Process based on file type
        if file.filename.endswith('.pdf'):
            text_content = extract_text_from_pdf(content)
        elif file.filename.endswith('.docx'):
            text_content = extract_text_from_docx(content)
        elif file.filename.endswith('.txt'):
            text_content = content.decode('utf-8')
        else:
            raise HTTPException(
        )
        
        # Update document status in the background
        background_tasks.add_task(
            document_service._process_document_async,
            document_id=document.id,
            db=db
        )
        
        return document
        
    except ValueError as e:
        logger.error(f"Validation error in document upload: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        db.rollback()
        logger.error(f"Error uploading document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to process document"
        )


@router.post("/batch/upload", response_model=List[DocumentResponse], status_code=status.HTTP_202_ACCEPTED)
async def batch_upload_documents(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    batch_data: DocumentUpload = Depends(),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> List[DocumentResponse]:
    """
    Upload and process multiple medical documents in a batch.
    
    Args:
        background_tasks: Background tasks manager
        files: List of uploaded document files
        batch_data: Batch metadata
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        List of accepted document information
    """
    try:
        # Check user's document quota
        from app.core.config import settings
        
        # Admin users have no quota
        if not current_user.is_admin:
            doc_count = db.query(MedicalDocument).filter(
                MedicalDocument.uploaded_by == current_user.id
            ).count()
            
            if doc_count + len(files) > settings.max_documents_per_user:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail=f"Maximum document limit of {settings.max_documents_per_user} reached"
                )
        
        # Validate file types
        for file in files:
            if not any(file.filename.lower().endswith(ext) for ext in settings.allowed_file_types):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File type not allowed. Allowed types: {settings.allowed_file_types}"
                )
        
        # Process each file
        documents = []
        for file in files:
            # Read file content
            content = await file.read()
            
            # Process based on file type
            if file.filename.endswith('.pdf'):
                text_content = extract_text_from_pdf(content)
            elif file.filename.endswith('.docx'):
                text_content = extract_text_from_docx(content)
            elif file.filename.endswith('.txt'):
                text_content = content.decode('utf-8')
            else:
                raise HTTPException(
            )
            
            # Create document
            document = MedicalDocument(
                title=file.filename,
                content=text_content,
                uploaded_by=current_user.id,
                metadata={"batch_id": batch_data.batch_id}
            )
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Update document status in the background
            background_tasks.add_task(
                document_service._process_document_async,
                document_id=document.id,
                db=db
            )
            
            documents.append(DocumentResponse.from_orm(document))
        
        return documents
        
    except ValueError as e:
        logger.error(f"Validation error in batch document upload: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        db.rollback()
        logger.error(f"Error uploading batch documents: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to process batch documents"
        )


@router.get("/", response_model=PaginatedResponse)
async def get_documents(
    pagination: PaginationParams = Depends(),
    document_type: Optional[DocumentType] = None,
    status: Optional[str] = None,
    search: Optional[str] = None,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> PaginatedResponse:
    """
    Get medical documents with pagination and filtering.
    
    Args:
        pagination: Pagination parameters
        document_type: Optional document type filter
        status: Optional document status filter
        search: Optional search query
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        Paginated list of documents
    """
    try:
        # Build query
        query = db.query(MedicalDocument)
        
        # Apply filters
        if document_type:
            query = query.filter(MedicalDocument.document_type == document_type)
        
        # Filter by status if provided
        if status:
            query = query.filter(MedicalDocument.status == status)
        
        # Non-admin users can only see their own documents
        if not current_user.is_admin:
            query = query.filter(MedicalDocument.uploaded_by == current_user.id)
        
        # Apply search
        if search:
            search = f"%{search}%"
            query = query.filter(
                (MedicalDocument.title.ilike(search)) |
                (MedicalDocument.description.ilike(search)) |
                (MedicalDocument.content.ilike(search))
            )
        
        # Get total count
        total = query.count()
        
        # Apply pagination
        documents = query.order_by(
            MedicalDocument.updated_at.desc()
        ).offset(
            pagination.offset
        ).limit(
            pagination.limit
        ).all()
        
        document_responses = [DocumentResponse.from_orm(doc) for doc in documents]
        
        return PaginatedResponse(
            items=document_responses,
            total=total,
            page=pagination.page,
            size=pagination.size
        )
        
    except Exception as e:
        logger.error(f"Failed to get documents: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve documents"
        )


@router.get("/stats", response_model=DocumentStatistics)
async def get_document_statistics(
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> JSONResponse:
    """
    Get comprehensive document statistics including counts by type, status, and storage usage.
    
    Args:
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        JSON response with document statistics
    """
    try:
        # Base query with access control
        base_query = db.query(MedicalDocument)
        if not current_user.is_admin:
            base_query = base_query.filter(MedicalDocument.uploaded_by == current_user.id)
        
        # Get total document count
        total_documents = base_query.count()
        
        # Get document count by type
        type_stats = base_query.with_entities(
            MedicalDocument.document_type,
            func.count(MedicalDocument.id).label('count')
        ).group_by(MedicalDocument.document_type).all()
        
        # Get document count by status
        status_stats = base_query.with_entities(
            MedicalDocument.status,
            func.count(MedicalDocument.id).label('count')
        ).group_by(MedicalDocument.status).all()
        
        # Calculate storage usage (in bytes)
        storage_usage = db.query(
            func.sum(func.length(MedicalDocument.content))
        ).scalar() or 0
        
        # Get recent activity
        recent_activity = base_query.order_by(
            MedicalDocument.updated_at.desc()
        ).limit(5).all()
        
        # Get document type distribution
        documents_by_type = [
            {"type": doc_type, "count": count} 
            for doc_type, count in type_stats
        ]
        
        # Get status distribution
        documents_by_status = [
            {"status": status, "count": count} 
            for status, count in status_stats
        ]
        
        # Get recent documents
        recent_docs = [
            {
                "id": str(doc.id),
                "title": doc.title,
                "type": doc.document_type.value,
                "status": doc.status.value,
                "updated_at": doc.updated_at.isoformat()
            }
            for doc in recent_activity
        ]
        
        stats = {
            "total_documents": total_documents,
            "documents_by_type": documents_by_type,
            "documents_by_status": documents_by_status,
            "storage_usage": {
                "bytes": storage_usage,
                "mb": round(storage_usage / (1024 * 1024), 2),
                "gb": round(storage_usage / (1024 * 1024 * 1024), 2)
            },
            "recent_activity": recent_docs
        }
        
        return JSONResponse(content=stats)
        
    except Exception as e:
        logger.error(f"Error getting document statistics: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve document statistics"
        )


@router.get("/batch/status/{batch_id}", response_model=Dict[str, Any])
async def get_batch_status(
    batch_id: str,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> Dict[str, Any]:
    """
    Get the status of a batch document upload.
    
    Args:
        batch_id: Batch ID to check status for
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        Batch status information
    """
    try:
        # Query documents with the batch ID
        query = db.query(MedicalDocument).filter(
            MedicalDocument.metadata["batch_id"].astext == batch_id
        )
        
        # Non-admin users can only see their own documents
        if not current_user.is_admin:
            query = query.filter(MedicalDocument.uploaded_by == current_user.id)
        
        documents = query.all()
        
        if not documents:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Batch not found or access denied"
            )
        
        # Calculate status counts
        status_counts = {}
        for doc in documents:
            status_counts[doc.status] = status_counts.get(doc.status, 0) + 1
        
        return {
            "batch_id": batch_id,
            "total_documents": len(documents),
            "status_counts": status_counts,
            "documents": [{
                "id": str(doc.id),
                "title": doc.title,
                "status": doc.status,
                "created_at": doc.created_at.isoformat(),
                "processed_at": doc.processed_at.isoformat() if doc.processed_at else None,
                "error": doc.error_message
            } for doc in documents]
        }
        
    except Exception as e:
        logger.error(f"Error getting batch status: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to get batch status"
        )


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: UUID,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> DocumentResponse:
    """
    Get a specific document by ID.
    
    Args:
        document_id: Document ID
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        Document information
    """
    try:
        document = db.query(MedicalDocument).filter(
            MedicalDocument.id == document_id
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        return DocumentResponse.from_orm(document)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve document"
        )


@router.post("/search", response_model=List[DocumentSearchResult])
async def search_documents(
    search_data: DocumentSearch,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
) -> List[DocumentSearchResult]:
    """
    Search medical documents using hybrid search (semantic + keyword).
    
    Args:
        search_data: Search parameters including query, filters, and pagination
        current_user: Current authenticated user
        db: Database session
        
    Returns:
        List of matching documents with relevance scores and highlights
    """
    try:
        # Base query with access control
        query = db.query(MedicalDocument)
        if not current_user.is_admin:
            query = query.filter(MedicalDocument.uploaded_by == current_user.id)
        
        # Initialize results list
        results = []
        
        # If we have a search query, perform hybrid search
        if search_data.query:
            # Perform semantic search if enabled
            if settings.enable_semantic_search:
                try:
                    # Get semantic search results from vector store
                    semantic_results = vector_service.search_similar(
                        query_text=search_data.query,
                        top_k=search_data.limit or 20,
                        namespace="medical-docs"
                    )
                    
                    # Get document IDs from semantic search
                    semantic_doc_ids = {result['metadata']['document_id'] for result in semantic_results}
                    semantic_scores = {result['metadata']['document_id']: result['score'] for result in semantic_results}
                    
                    # Get documents from database
                    docs = query.filter(MedicalDocument.id.in_([UUID(doc_id) for doc_id in semantic_doc_ids])).all()
                    
                    # Create search results with semantic scores
                    for doc in docs:
                        doc_id = str(doc.id)
                        results.append(DocumentSearchResult(
                            document=DocumentResponse.from_orm(doc),
                            score=semantic_scores.get(doc_id, 0.0),
                            highlights={
                                "content": [search_data.query],  # Simple highlight for now
                                "title": [search_data.query] if search_data.query.lower() in doc.title.lower() else []
                            }
                        ))
                    
                    # If we have enough results, return them
                    if len(results) >= (search_data.limit or 20):
                        return results[:search_data.limit] if search_data.limit else results
                        
                except Exception as e:
                    logger.warning(f"Semantic search failed, falling back to keyword search: {str(e)}")
            
            # Fall back to keyword search
            search_terms = f"%{search_data.query}%"
            keyword_conditions = [
                MedicalDocument.title.ilike(search_terms),
                MedicalDocument.content.ilike(search_terms),
                MedicalDocument.metadata['keywords'].astext.ilike(search_terms)
            ]
            
            # Get keyword search results
            keyword_docs = query.filter(or_(*keyword_conditions)).all()
            
            # Add keyword results that aren't already in results
            existing_ids = {str(r.document.id) for r in results}
            for doc in keyword_docs:
                if str(doc.id) not in existing_ids:
                    results.append(DocumentSearchResult(
                        document=DocumentResponse.from_orm(doc),
                        score=0.5,  # Lower score for keyword matches
                        highlights={
                            "content": [search_data.query],
                            "title": [search_data.query] if search_data.query.lower() in doc.title.lower() else []
                        }
                    ))
        else:
            # No search query, just apply filters and sorting
            docs = query.all()
            results = [
                DocumentSearchResult(
                    document=DocumentResponse.from_orm(doc),
                    score=0.0
                )
                for doc in docs
            ]
        
        # Apply filters
        if search_data.document_type:
            results = [r for r in results if r.document.document_type == search_data.document_type]
            
        if search_data.status:
            results = [r for r in results if r.document.status == search_data.status]
        
        # Apply date range filter
        if search_data.start_date:
            results = [r for r in results if r.document.created_at >= search_data.start_date]
        if search_data.end_date:
            results = [r for r in results if r.document.created_at <= search_data.end_date]
        
        # Sort results by score (descending) and then by the requested field
        sort_field = search_data.sort_by if search_data.sort_by != 'score' else None
        reverse_sort = search_data.sort_order == "desc"
        
        def get_sort_key(result):
            key = []
            # Primary sort by score (descending)
            key.append(-result.score if result.score is not None else 0)
            
            # Secondary sort by requested field
            if sort_field and hasattr(result.document, sort_field):
                field_value = getattr(result.document, sort_field)
                # Handle None values for proper sorting
                key.append(field_value is None)
                key.append(field_value)
            
            # Tertiary sort by document ID for stability
            key.append(result.document.id)
            return tuple(key)
        
        results.sort(key=get_sort_key, reverse=reverse_sort)
        
        # Apply pagination
        start = search_data.offset or 0
        end = start + (search_data.limit or len(results)) if search_data.limit else len(results)
        
        return results[start:end]
        
    except Exception as e:
        logger.error(f"Error searching documents: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to search documents"
        )


@router.delete("/{document_id}")
async def delete_document(
    document_id: UUID,
    admin_user = Depends(get_admin_user),
    db: Session = Depends(get_db)
):
    """
    Delete a medical document (admin only).
    
    Args:
        document_id: Document ID
        admin_user: Admin user
        db: Database session
        
    Returns:
        Success message
    """
    try:
        document = db.query(MedicalDocument).filter(
            MedicalDocument.id == document_id
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        # Delete from vector database
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).all()
        
        if chunks:
            vector_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]
            if vector_ids:
                vector_service.delete_vectors(vector_ids)
        
        # Delete from database
        db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).delete()
        
        db.delete(document)
        db.commit()
        
        logger.info(f"Document deleted: {document_id}")
        return {"message": "Document deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to delete document"
        )


@router.post("/{document_id}/reprocess")
async def reprocess_document(
    document_id: UUID,
    admin_user = Depends(get_admin_user),
    db: Session = Depends(get_db)
):
    """
    Reprocess a document for vector embeddings (admin only).
    
    Args:
        document_id: Document ID
        admin_user: Admin user
        db: Database session
        
    Returns:
        Success message
    """
    try:
        document = db.query(MedicalDocument).filter(
            MedicalDocument.id == document_id
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        # Delete existing chunks and vectors
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).all()
        
        if chunks:
            vector_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]
            if vector_ids:
                vector_service.delete_vectors(vector_ids)
        
        db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).delete()
        
        # Mark as unprocessed
        document.processed = False
        db.commit()
        
        # Process document
        await process_document_async(document_id, db)
        
        logger.info(f"Document queued for reprocessing: {document_id}")
        return {"message": "Document queued for reprocessing"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to reprocess document: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to reprocess document"
        )


# Helper functions
def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        import PyPDF2
        from io import BytesIO
        
        pdf_reader = PyPDF2.PdfReader(BytesIO(content))
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text()
        
        return text
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from PDF"
        )


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        from io import BytesIO
        
        doc = Document(BytesIO(content))
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from DOCX"
        )


async def process_document_async(document_id: UUID, db: Session):
    """Process document for vector embeddings."""
    try:
        document = db.query(MedicalDocument).filter(
            MedicalDocument.id == document_id
        ).first()
        
        if not document:
            logger.error(f"Document not found for processing: {document_id}")
            return
        
        # Chunk the document
        chunks = ai_service.chunk_medical_document(document.content)
        
        if not chunks:
            logger.error(f"No chunks created for document: {document_id}")
            return
        
        # Create document metadata
        document_metadata = {
            "title": document.title,
            "document_type": document.document_type.value,
            "source": document.source,
            "authors": document.authors,
            "keywords": document.keywords,
            "publication_date": document.publication_date.isoformat() if document.publication_date else None
        }
        
        # Process chunks for vector storage
        vectors = vector_service.process_document_for_vectors(
            document_id=str(document_id),
            chunks=chunks,
            document_metadata=document_metadata
        )
        
        if not vectors:
            logger.error(f"No vectors created for document: {document_id}")
            return
        
        # Store vectors in Pinecone
        success = vector_service.upsert_vectors(vectors)
        
        if not success:
            logger.error(f"Failed to store vectors for document: {document_id}")
            return
        
        # Create document chunks in database
        for i, chunk in enumerate(chunks):
            chunk_record = DocumentChunk(
                document_id=document_id,
                content=chunk["content"],
                chunk_index=i,
                vector_id=f"{document_id}_{i}",
                start_position=chunk.get("start_position", 0),
                end_position=chunk.get("end_position", 0)
            )
            db.add(chunk_record)
        
        # Mark document as processed
        document.processed = True
        db.commit()
        
        logger.info(f"Document processed successfully: {document_id}")
        
    except Exception as e:
        logger.error(f"Failed to process document {document_id}: {e}")
        db.rollback()


@router.get("/stats/summary")
async def get_document_stats(
    admin_user = Depends(get_admin_user),
    db: Session = Depends(get_db)
):
    """
    Get document collection statistics (admin only).
    
    Args:
        admin_user: Admin user
        db: Database session
        
    Returns:
        Document statistics
    """
    try:
        # Total documents
        total_documents = db.query(MedicalDocument).count()
        
        # Processed documents
        processed_documents = db.query(MedicalDocument).filter(
            MedicalDocument.processed == True
        ).count()
        
        # Documents by type
        document_types = db.query(
            MedicalDocument.document_type,
            db.func.count(MedicalDocument.id)
        ).group_by(MedicalDocument.document_type).all()
        
        # Total chunks
        total_chunks = db.query(DocumentChunk).count()
        
        # Vector database stats
        vector_stats = vector_service.get_index_stats()
        
        return {
            "total_documents": total_documents,
            "processed_documents": processed_documents,
            "processing_rate": processed_documents / total_documents if total_documents > 0 else 0,
            "document_types": {str(doc_type): count for doc_type, count in document_types},
            "total_chunks": total_chunks,
            "vector_database": vector_stats
        }
        
    except Exception as e:
        logger.error(f"Failed to get document stats: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve document statistics"
        )
