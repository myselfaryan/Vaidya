"""
Document processing service for medical documents.
"""

import os
import uuid
from typing import List, Dict, Any, Optional, Tuple, BinaryIO
from datetime import datetime
import io

from loguru import logger
from sqlalchemy.orm import Session
import PyPDF2
from docx import Document

from ..models.models import MedicalDocument, DocumentChunk, DocumentStatus, DocumentType
from ..models.schemas import DocumentUpload, DocumentResponse, DocumentChunkResponse
from ..core.config import settings
from .vector_service import vector_service
from ..utils.document_processor import DocumentProcessor


class DocumentService:
    """Service for document processing and management."""
    
    def __init__(self):
        """Initialize the document service."""
        self.document_processor = DocumentProcessor(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap
        )
    
    async def process_uploaded_file(
        self,
        file: BinaryIO,
        filename: str,
        document_data: DocumentUpload,
        user_id: uuid.UUID,
        db: Session
    ) -> DocumentResponse:
        """
        Process an uploaded document file.
        
        Args:
            file: Uploaded file object
            filename: Original filename
            document_data: Document metadata
            user_id: ID of the user who uploaded the document
            db: Database session
            
        Returns:
            Processed document information
        """
        try:
            # Read file content
            content = file.read()
            
            # Validate file size
            if len(content) > settings.max_file_size:
                raise ValueError(f"File size exceeds maximum allowed size of {settings.max_file_size} bytes")
            
            # Extract text based on file type
            text = self._extract_text_from_file(content, filename)
            
            # Create document record
            db_document = MedicalDocument(
                title=document_data.title,
                description=document_data.description,
                document_type=document_data.document_type,
                source=document_data.source,
                content=text[:settings.max_text_length],
                status=DocumentStatus.PROCESSING,
                uploaded_by=user_id,
                metadata=document_data.metadata or {}
            )
            db.add(db_document)
            db.commit()
            db.refresh(db_document)
            
            # Process document asynchronously
            await self._process_document_async(db_document.id, db)
            
            return DocumentResponse.from_orm(db_document)
            
        except Exception as e:
            logger.error(f"Failed to process uploaded file: {e}")
            db.rollback()
            raise
    
    def _extract_text_from_file(self, content: bytes, filename: str) -> str:
        """
        Extract text from file content based on file type.
        
        Args:
            content: File content as bytes
            filename: Original filename
            
        Returns:
            Extracted text
        """
        try:
            if filename.lower().endswith('.pdf'):
                return self._extract_text_from_pdf(content)
            elif filename.lower().endswith(('.docx', '.doc')):
                return self._extract_text_from_docx(content)
            elif filename.lower().endswith(('.txt', '.md')):
                return content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {filename.split('.')[-1]}")
        except Exception as e:
            logger.error(f"Failed to extract text from file: {e}")
            raise ValueError(f"Failed to process file: {str(e)}")
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        try:
            doc = Document(io.BytesIO(content))
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise
    
    async def _process_document_async(
        self,
        document_id: uuid.UUID,
        db: Session
    ) -> None:
        """
        Process document asynchronously for vector embeddings.
        
        Args:
            document_id: ID of the document to process
            db: Database session
        """
        try:
            # Get document from database
            document = db.query(MedicalDocument).filter(
                MedicalDocument.id == document_id
            ).first()
            
            if not document:
                logger.error(f"Document not found: {document_id}")
                return
            
            # Update status to processing
            document.status = DocumentStatus.PROCESSING
            document.processed_at = datetime.utcnow()
            db.commit()
            
            try:
                # Process document into chunks
                chunks = self.document_processor.process_document(
                    document_id=str(document.id),
                    text=document.content,
                    document_metadata={
                        "title": document.title,
                        "source": document.source,
                        "document_type": document.document_type.value,
                        **document.metadata
                    }
                )
                
                # Generate embeddings for chunks
                chunk_texts = [chunk['text'] for chunk in chunks]
                chunk_embeddings = vector_service.generate_embeddings(chunk_texts)
                
                # Prepare vectors for Pinecone
                vectors = []
                for i, (chunk, embedding) in enumerate(zip(chunks, chunk_embeddings)):
                    vector_id = f"{document_id}_{i}"
                    chunk_metadata = chunk['metadata']
                    
                    # Save chunk to database
                    db_chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=i,
                        content=chunk['text'],
                        metadata=chunk_metadata,
                        vector_id=vector_id
                    )
                    db.add(db_chunk)
                    
                    # Prepare vector for Pinecone
                    vectors.append({
                        'id': vector_id,
                        'values': embedding,
                        'metadata': {
                            **chunk_metadata,
                            'text': chunk['text'],
                            'document_id': str(document_id),
                            'processed_at': datetime.utcnow().isoformat()
                        }
                    })
                
                # Save all chunks to database
                db.commit()
                
                # Upsert vectors to Pinecone
                if vectors:
                    vector_service.upsert_vectors(vectors, namespace="medical-docs")
                
                # Update document status
                document.status = DocumentStatus.PROCESSED
                document.chunk_count = len(chunks)
                document.processed_at = datetime.utcnow()
                db.commit()
                
                logger.info(f"Successfully processed document: {document_id} with {len(chunks)} chunks")
                
            except Exception as e:
                logger.error(f"Failed to process document {document_id}: {e}")
                document.status = DocumentStatus.FAILED
                document.error_message = str(e)
                db.commit()
                raise
                
        except Exception as e:
            logger.error(f"Error in async document processing: {e}")
            raise


# Global document service instance
document_service = DocumentService()
